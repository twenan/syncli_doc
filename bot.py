import asyncio
import os
import datetime
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from fpdf import FPDF
from datetime import datetime, timedelta
from num2words import num2words
from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.storage.memory import MemoryStorage
from config import Config, load_config

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Загрузка конфигурации бота
config: Config = load_config()
BOT_TOKEN: str = config.tg_bot.token

# Создание бота и диспетчера
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())

# Определение состояний для FSM
class ContractStates(StatesGroup):
    GET_CUSTOMER_NAME = State()
    GET_CONTRACT_AMOUNT = State()
    GET_PRODUCT_NAME = State()
    GET_BANK_DETAILS = State()

# Шаблон договора
TEMPLATE_PATH = "template.docx"


def debug_placeholders(doc, placeholders):
    logging.info("🔍 **Начинаем диагностику плейсхолдеров**")
    found = set()

    for paragraph in doc.paragraphs:
        for key in placeholders.keys():
            if key in paragraph.text:
                logging.info(f"✅ Найден плейсхолдер ВНЕ таблицы: '{key}' в тексте: {paragraph.text}")
                found.add(key)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in placeholders.keys():
                        if key in paragraph.text:
                            logging.info(f"✅ Найден плейсхолдер В ТАБЛИЦЕ: '{key}' в ячейке: {paragraph.text}")
                            found.add(key)

    if not found:
        logging.warning("⚠️ **Плейсхолдеры вообще не найдены в документе!**")

    return found


# ✅ Функция для замены плейсхолдеров
def replace_placeholders(doc, placeholders):
    replaced = set()  # Множество замененных плейсхолдеров

    # 🔹 Обработка плейсхолдеров в параграфах
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)  # Полный текст из runs
        logging.info(f"📌 Текст параграфа ДО замены: {full_text}")

        modified_text = full_text  # Создаём копию текста для работы

        # 🔹 Проходим по каждому плейсхолдеру и заменяем
        for key, value in placeholders.items():
            if key.lower() in modified_text.lower():  # Проверяем наличие в тексте
                logging.info(f"🔄 Заменяем '{key}' на '{value}'")

                # 🔹 Заменяем все плейсхолдеры, если они идут подряд
                modified_text = modified_text.replace(key, value)

        # 🔹 Очищаем текущие runs перед обновлением текста
        for run in paragraph.runs:
            run.text = ""

        # 🔹 Записываем обновленный текст в первый run
        if paragraph.runs:
            paragraph.runs[0].text = modified_text

        # 🔹 Форматирование
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(13)

        # 🔹 Выравнивание
        if "{сегодняшняя дата 1}" in modified_text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 🔹 Обработка плейсхолдеров в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    modified_text = full_text  # Копия текста

                    # 🔹 Проходим по каждому плейсхолдеру
                    for key, value in placeholders.items():
                        if key.lower() in modified_text.lower():
                            logging.info(f"🔄 Заменяем '{key}' в таблице на '{value}'")
                            modified_text = modified_text.replace(key, value)

                    # 🔹 Очищаем текущие runs перед обновлением текста
                    for run in paragraph.runs:
                        run.text = ""

                    # 🔹 Записываем обновленный текст в первый run
                    if paragraph.runs:
                        paragraph.runs[0].text = modified_text  
                        
                    # Форматируем текст
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(13)
                    

# ✅ Функция для создания PDF из DOCX
def create_pdf(docx_path, pdf_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
    pdf.set_font("DejaVu", size=13)

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        pdf.multi_cell(0, 10, paragraph.text)

    pdf.output(pdf_path)

# ✅ Хендлер команды /start
@dp.message(Command("start"))
async def start(message: types.Message, state: FSMContext):
    welcome_text = (
        "🤖 **Что умеет этот бот?**\n\n"
        "Этот бот предназначен для помощи в заполнении договоров.\n\n"
        "Соберём все данные для заполнения договора шаг за шагом. Готовы начать?"
    )

    keyboard = types.ReplyKeyboardMarkup(
        keyboard=[
            [types.KeyboardButton(text="🚀 Начать заполнение договора")]
        ],
        resize_keyboard=True
    )

    await message.answer(welcome_text, reply_markup=keyboard, parse_mode="Markdown")

# ✅ Хендлер для кнопки "🚀 Начать заполнение договора"
@dp.message(lambda message: message.text == "🚀 Начать заполнение договора")
async def start_contract_filling(message: types.Message, state: FSMContext):
    await message.answer("Введите ФИО заказчика:")
    await state.set_state(ContractStates.GET_CUSTOMER_NAME)

@dp.message(ContractStates.GET_CUSTOMER_NAME)
async def get_customer_name(message: types.Message, state: FSMContext):
    await state.update_data(customer_name=message.text)
    await message.answer("Введите сумму договора (цифрами):")
    await state.update_data(customer_name=message.text)  # ✅ Сохраняем в state
    await state.set_state(ContractStates.GET_CONTRACT_AMOUNT)  # 🔹 Указываем следующее состояние

@dp.message(ContractStates.GET_CONTRACT_AMOUNT)
async def get_contract_amount(message: types.Message, state: FSMContext):
    await state.update_data(contract_amount=message.text.strip())
    await message.answer("Введите название товара в родительном падеже:")
    await state.set_state(ContractStates.GET_PRODUCT_NAME)

@dp.message(ContractStates.GET_PRODUCT_NAME)
async def get_product_name(message: types.Message, state: FSMContext):
    await state.update_data(product_name=message.text)
    await message.answer("Введите банковские реквизиты:")
    await state.set_state(ContractStates.GET_BANK_DETAILS)


import re  # Добавляем регулярные выражения

# Парсер реквизитов заказчика
def parse_bank_details(raw_text):
    """Разбирает текст и выделяет из него реквизиты клиента."""
    
    patterns = {
               "customer_name": r"(?:Наименование|ФИО|ИП|Индивидуальный предприниматель):?\s*([А-ЯЁа-яё\s]+)",
        "inn": r"ИНН:?\s*(\d{10,12})",  # Берем первый найденный ИНН (основной)
        "ogrnip": r"ОГРНИП:?\s*(\d+)",
        "account_number": r"(?:р\.с|р\.с\.|р/с|расс?\.\s?сч[её]т|расс?\s?сч[её]т|рас\.сч[её]т|рассч[её]тный\s?сч[её]т|расч[её]тный\s?сч[её]т):?\s*(\d{20})",
        "bank_name": r"(?:банк|наименование банка):?\s*([А-ЯЁа-яё\s]+(?:банк|БАНК|Bank)?)",
        "bik": r"(?:БИК\s?банка|БИК):?\s*(\d{9})",
        "correspondent_account": r"(?:корр?сч[её]т|кор\.?\s?сч[её]т|корр?\s?сч[её]т|корреспондентский\s?сч[её]т|к\.с\.|к\.с|к/с):?\s*(\d{20})",
        "kpp": r"КПП:?\s*(\d{9})",
        "okpo": r"ОКПО:?\s*(\d{8})",  # ОКПО (обычно 8 цифр)
        "oktmo": r"ОКТМО:?\s*(\d{8})",  # ОКТМО (обычно 8 цифр)
        "phone": r"Тел(?:ефон)?:?\s*([\d\-\+\(\)\s]{10,16})"  # Извлечение телефона
    }
    extracted_data = {}

    for field, pattern in patterns.items():
        matches = re.findall(pattern, raw_text, re.IGNORECASE)
        
        # Берем первый найденный элемент для ИНН, а для остальных последний найденный
        if field == "inn":
            extracted_data[field] = matches[0] if matches else "Не указано"
        else:
            extracted_data[field] = matches[-1] if matches else "Не указано"
    
    return extracted_data


# ✅ Обработчик ввода банковских реквизитов
@dp.message(ContractStates.GET_BANK_DETAILS)
async def get_bank_details(message: types.Message, state: FSMContext):
    raw_text = message.text.strip()  # Убираем лишние пробелы

    # Парсим данные из сообщения
    parsed_data = parse_bank_details(raw_text)

    # Получаем сохраненные данные
    data = await state.get_data()

    today_date = datetime.now().strftime("%d.%m.%Y")
    future_date = (datetime.now() + timedelta(days=45)).strftime("%d.%m.%Y")

    # Проверяем корректность ввода суммы
    try:
        contract_amount = int(data.get("contract_amount", "0").replace(" ", ""))
        logging.info(f"💰 Сумма работ: {contract_amount}")
    except ValueError:
        logging.error(f"❌ Некорректное значение суммы работ: {data.get('contract_amount', '0')}")
        contract_amount = 0

    # Заполняем placeholders из `parsed_data`
    placeholders = {
        "{сегодняшняя дата 1}": today_date,
        "{заказчик 1}": f"Индивидуальный Предприниматель {parsed_data['customer_name']}",
        "{название товара в родительном падеже}": data.get("product_name", "Пустое значение"),
        "{сегодняшняя дата}": today_date,
        "{полтора месяца вперед от сегодняшней даты}": future_date,
        "{contract_amount}": str(contract_amount),
        "{стоимость работ прописью}": num2words(contract_amount, lang="ru"),
        "{юридический адрес заказчика}": parsed_data["customer_name"],
        "{ИНН заказчика}": parsed_data["inn"],
        "{ОГРНИП заказчика}": parsed_data["ogrnip"],
        "{ОКПО заказчика}": parsed_data["okpo"], 
        "{ОКТМО заказчика}": parsed_data["oktmo"],
        "{расчетный счет заказчика}": parsed_data["account_number"],
        "{банк заказчика}": parsed_data["bank_name"],
        "{корреспондентский счет банка заказчика}": parsed_data["correspondent_account"],
        "{БИК банка заказчика}": parsed_data["bik"],
        "{телефон заказчика}": parsed_data["phone"]
    }

    # Проверяем, какие поля не удалось заполнить
    missing_fields = [key for key, value in placeholders.items() if value == "Не указано"]

    if missing_fields:
        missing_text = "\n".join(f"🔹 {field}" for field in missing_fields)
        await message.answer(f"⚠️ Не удалось автоматически распознать следующие поля:\n{missing_text}\n\nВведите их вручную.")
        return

    # Логируем значения для отладки
    logging.info("📌 Передаем значения для заполнения:")
    for key, value in placeholders.items():
        logging.info(f"{key}: {value}")

    # ✅ Генерация уникального имени файла
    customer_name = parsed_data["customer_name"].replace(" ", "_")
    file_date = datetime.now().strftime("%d-%m-%Y")  # Дата в имени файла
    file_name = f"Договор_{customer_name}_{file_date}"

    # ✅ Пути к файлам
    docx_output_path = f"/home/anna/syncli_doc/syncli_doc/{file_name}.docx"
    pdf_output_path = f"/home/anna/syncli_doc/syncli_doc/{file_name}.pdf"

    # ✅ Загрузка шаблона и замена плейсхолдеров
    doc = Document(TEMPLATE_PATH)
    replace_placeholders(doc, placeholders)

    # ✅ Сохранение DOCX
    try:
        doc.save(docx_output_path)
    except Exception as e:
        logging.error(f"❌ Ошибка при сохранении DOCX: {str(e)}")
        await message.answer(f"Ошибка при сохранении DOCX: {str(e)}")
        return

    # ✅ Создание PDF
    try:
        create_pdf(docx_output_path, pdf_output_path)
    except Exception as e:
        logging.error(f"❌ Ошибка при создании PDF: {str(e)}")
        await message.answer(f"Ошибка при создании PDF: {str(e)}")
        return

    # ✅ Отправка готового договора
    if os.path.exists(docx_output_path) and os.path.exists(pdf_output_path):
        try:
            await message.answer(f"✅ Готовый договор `{file_name}` создан и отправлен:")
            await message.answer_document(types.FSInputFile(docx_output_path))
            await message.answer_document(types.FSInputFile(pdf_output_path))
        except Exception as e:
            logging.error(f"❌ Ошибка при отправке файла: {str(e)}")
            await message.answer(f"Ошибка при отправке файла: {str(e)}")
    else:
        logging.error("❌ Ошибка: файлы не были созданы.")
        await message.answer("Ошибка: не удалось создать файлы договора.")

    await state.clear()


# ✅ Основная функция запуска бота
async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
